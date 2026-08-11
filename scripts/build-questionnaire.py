# Builds docs/portfolio-v3-content-questionnaire.docx
# One question per former amber chip on the live v3 pages, plus the items
# that are now missing from the page entirely because the chip was the only
# thing in them.
import io, os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

INK = RGBColor(0x10, 0x18, 0x28)
BODY = RGBColor(0x3C, 0x4A, 0x63)
MUTED = RGBColor(0x5B, 0x6D, 0x84)
BRAND = RGBColor(0x2C, 0x46, 0xA8)
CORAL = RGBColor(0xC2, 0x3A, 0x63)

doc = Document()
st = doc.styles['Normal']
st.font.name = 'Aptos'
st.font.size = Pt(10.5)
st.font.color.rgb = BODY
st.paragraph_format.space_after = Pt(6)
st.paragraph_format.line_spacing = 1.15
for s in doc.sections:
    s.top_margin = s.bottom_margin = Cm(2.2)
    s.left_margin = s.right_margin = Cm(2.4)


def para(text='', size=10.5, bold=False, italic=False, color=BODY,
         before=0, after=6, align=None, style=None):
    p = doc.add_paragraph(style=style)
    r = p.add_run(text)
    r.bold, r.italic = bold, italic
    r.font.size = Pt(size)
    r.font.color.rgb = color
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    if align is not None:
        p.alignment = align
    return p


def rule(after=10):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(after)
    pPr = p._p.get_or_add_pPr()
    bd = OxmlElement('w:pBdr')
    bt = OxmlElement('w:bottom')
    bt.set(qn('w:val'), 'single'); bt.set(qn('w:sz'), '6')
    bt.set(qn('w:space'), '1'); bt.set(qn('w:color'), 'C9D5E8')
    bd.append(bt); pPr.append(bd)


def answer_box(lines=3):
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    t.autofit = True
    cell = t.cell(0, 0)
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F4F7FD')
    tcPr.append(shd)
    borders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        e = OxmlElement('w:' + edge)
        e.set(qn('w:val'), 'single'); e.set(qn('w:sz'), '6')
        e.set(qn('w:space'), '0'); e.set(qn('w:color'), 'C9D5E8')
        borders.append(e)
    tcPr.append(borders)
    cell.paragraphs[0].text = ''
    for i in range(lines):
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.3
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return t


def question(num, where, ask, hint, lines=3, flag=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(1)
    r = p.add_run('Q%s' % num); r.bold = True; r.font.size = Pt(10); r.font.color.rgb = BRAND
    r2 = p.add_run('   ' + where); r2.font.size = Pt(9); r2.font.color.rgb = MUTED
    para(ask, size=11, bold=True, color=INK, after=2)
    if flag:
        para(flag, size=9, italic=True, color=CORAL, after=3)
    para(hint, size=9.5, color=MUTED, after=4)
    answer_box(lines)


# ───────────────────────── cover ─────────────────────────
para('Portfolio content questionnaire', size=22, bold=True, color=INK, after=2)
para('The live v3 site: home, work, projects, experience. Everything only you can answer.',
     size=11.5, color=MUTED, after=12)

for k, v in [('Prepared for', 'Manfred Siew'),
             ('Covers', 'public/index.html, public/work.html, public/projects.html, public/experience.html'),
             ('Questions', '28, across 3 parts'),
             ('Date', '11 August 2026')]:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    r = p.add_run(k.ljust(14)); r.bold = True; r.font.size = Pt(9.5); r.font.color.rgb = INK
    r2 = p.add_run(v); r2.font.size = Pt(9.5); r2.font.color.rgb = BODY

rule(after=12)

para('How to use this', size=13, bold=True, color=INK, before=6, after=4)
for b in [
    'The amber "needs copy" chips have been taken off the site. This document is now the only '
    'place those gaps are recorded, so nothing gets lost by them no longer being visible.',
    'Answer in the shaded boxes, in plain language. The wording gets edited into the page afterwards, '
    'so a rough sentence is worth more than a polished blank.',
    '"I do not know" and "cannot publish that" are real answers. Write them in. A gap you flag is a '
    'gap that can be handled; a gap left silent reads to a stranger as a claim with nothing behind it.',
    'Part 1 is different from the rest: those four answers are for content that is currently missing '
    'from the page altogether, because the chip was the only thing standing in for it. Until they are '
    'answered, the page is short by one case study step and three project outcomes.',
]:
    p = doc.add_paragraph(style='List Bullet')
    r = p.add_run(b); r.font.size = Pt(10); r.font.color.rgb = BODY
    p.paragraph_format.space_after = Pt(4)

# ─────────────────── part 1: holes in the page ───────────────────
doc.add_page_break()
para('Part 1   ·   The four the page is now visibly short of', size=15, bold=True, color=INK, after=3)
para('Each of these had a chip and nothing else. Removing the chip removed the only content in the slot, '
     'so the slot itself came out. These four answers put content back on the page rather than replacing '
     'a marker.', size=10, color=MUTED, after=8)

question(1, 'work.html  ·  GRID  ·  step 6',
         'What would you do differently on GRID?',
         'One honest line. The other two case studies both answer this and GRID now stops at step 5, which '
         'reads as though the question was avoided. Something you would set up earlier, agree earlier, or '
         'not assume again is enough.',
         lines=4,
         flag='The whole step is currently missing from the page.')

question(2, 'projects.html  ·  Whitehacks 2025',
         'What was the outcome of Whitehacks 2025?',
         'Placed, did not place, or what was assessed. "Took part, did not place" is a fine answer and is '
         'better than the row having no outcome at all.',
         lines=2,
         flag='The Outcome block is currently missing from this row.')

question(3, 'projects.html  ·  Iron Viz Student Edition',
         'What was the outcome of Iron Viz Student Edition?',
         'Same shape as above.',
         lines=2,
         flag='The Outcome block is currently missing from this row.')

question(4, 'projects.html  ·  WorldSkills Training',
         'What was the outcome of the WorldSkills aircraft maintenance training?',
         'This one is a credential rather than a contest, so the useful answer may be what you came away '
         'certified or assessed on rather than a placing.',
         lines=2,
         flag='The Outcome block is currently missing from this row.')

# ─────────────────── part 2: work page ───────────────────
doc.add_page_break()
para('Part 2   ·   The three case studies', size=15, bold=True, color=INK, after=3)
para('These sit on work.html. The copy around them already reads as finished, so each answer here makes an '
     'existing paragraph more specific rather than filling an empty one.', size=10, color=MUTED, after=8)

para('GRID', size=12, bold=True, color=INK, before=8, after=2)

question(5, 'work.html  ·  GRID  ·  step 2',
         'On GRID, what did you consider and reject?',
         'The framing step currently only says what you chose. What was the other option, and why did it '
         'lose? For example, another intake route, another extraction approach, or letting finance chase it.',
         lines=3)

question(6, 'work.html  ·  GRID  ·  step 5',
         'Who is running the GRID pilot, and how big is it?',
         'A unit name, a user count, or a number of receipts processed. Any one of them turns "in pilot now" '
         'from an assertion into a fact.',
         lines=2)

para('MILES / MAVIS', size=12, bold=True, color=INK, before=10, after=2)
para('Three of these confirm wording that was rewritten on 11 August from what you told me: that you were '
     'not the solution owner, that validation ran through the product owner and the stakeholder groups, and '
     'that the lesson learnt is about interviewing several people per role.',
     size=9.5, italic=True, color=MUTED, after=6)

question(7, 'work.html  ·  MILES / MAVIS  ·  step 3',
         'Is the new role wording right, and who did hold the solution ownership?',
         'The page now reads: "Full-stack developer. I was not the solution owner on this one: scope sat with '
         'the product owner, and I owned the build." Correct it if the split was different, and say who the '
         'product owner or solution owner actually was if you want them named.',
         lines=4)

question(8, 'work.html  ·  MILES / MAVIS  ·  step 3',
         'What was the timeline, the team size, or the hard constraint on MILES / MAVIS?',
         'GRID says "a team of five, two months". This case study says nothing comparable, so it reads as the '
         'smaller piece of work when it is the larger one.',
         lines=3)

question(9, 'work.html  ·  MILES / MAVIS  ·  step 1',
         'Who hit this problem, and what was it costing them?',
         'The page now names the people running and maintaining the Air Specialist Vehicle fleet. What was the '
         'actual cost of the old way: time, a missed servicing, someone driving unqualified, a report that took '
         'a day to assemble?',
         lines=3)

question(10, 'work.html  ·  MILES / MAVIS  ·  step 2',
          'What did you decide the problem actually was, and what did you rule out?',
          'Same question as Q5, for this project. The framing step reads as a description of the solution rather '
          'than a decision you made.',
          lines=3)

question(11, 'work.html  ·  MILES / MAVIS  ·  steps 5 and 6',
          'Is the validation and lesson-learnt wording right?',
          'Validation now reads that the product owner and the stakeholder groups signed it off, and that the '
          'pain points went back to them to check they had been read correctly from the ground up. Step 6 now '
          'reads that you would interview several people in the same role, before build, and have them '
          'cross-check each other. Correct anything that overstates or understates it.',
          lines=5)

para('Vibe Coding Masterclass', size=12, bold=True, color=INK, before=10, after=2)

question(12, 'work.html  ·  Masterclass  ·  step 2',
          'On the masterclass, what did you consider and reject?',
          'For example, a different format, a different length, a technical audience instead of a mixed one.',
          lines=3)

question(13, 'work.html  ·  Masterclass  ·  step 3',
          'Which year was the 6 August run?',
          'The page says "Run on 6 August for a cohort of 170" with no year, which dates badly the moment '
          'someone reads it in a different one.',
          lines=1)

# ─────────────────── part 3: project log ───────────────────
doc.add_page_break()
para('Part 3   ·   The project log', size=15, bold=True, color=INK, after=3)
para('These sit on projects.html. Most are one line each. The pattern behind nearly all of them is the same: '
     'a number or an acronym is on the page with nothing behind it.', size=10, color=MUTED, after=8)

log = [
    ('FUEL (Aircraft Refuelling Strategy Planner)',
     'Which unit is piloting FUEL, and what is the actual forecasting stack?',
     'The page says "an AI-assisted forecasting model" and "currently a working pilot". Both are vague enough '
     'that a technical reader will ask. Name the model or the method if you can.', 3),
    ('MatFlow (Supply and Demand Logistics Pipeline)',
     'Which unit is piloting MatFlow, and what is the pipeline actually built on?',
     'Same shape as the FUEL question. "A custom automated request pipeline" does not say what runs it.', 3),
    ('RSAF Facility Booking App',
     'Who can confirm the two-day build, or what dates did it run between?',
     'Two days end-to-end is the most impressive claim in the log and the least supported. A name or a pair of '
     'dates fixes it.', 2),
    ('SSB Loan Tracking System',
     'What does SSB stand for, and what do they do?',
     'The row uses the acronym twice and never expands it.', 2),
    ('Workplace Check In/Out App',
     'How was the efficiency and accuracy improvement measured, and who reported it?',
     'The page says "significantly improved". If it was a felt improvement rather than a measured one, say so '
     'and the wording gets softened to match.', 3),
    ('App Development for Poly Forum 2024 (SYLP)',
     'Where does the 500-student figure come from, and who counted it?',
     'Registered, attended, or capacity. They are three different numbers.', 2),
    ('R&D for Vibe-Coding Code Apps',
     'Who measured "months to days", and against which project?',
     'This is the strongest efficiency claim on the page. Against what baseline, and whose measurement?', 3),
    ('App Marketing Skill (marketing-pr)',
     'Who has installed or used it, and what did a user hit that you had not designed for?',
     'Anything real: a colleague, a GitHub star, a bug someone reported. It is an open-source release, so '
     'evidence of anyone else touching it counts for a lot.', 3),
    ('Google AI Studio System Prompt',
     'What incident made you write it, and roughly how many people have you handed it to?',
     'The specific moment an assistant went off and did something unasked is the story here.', 3),
    ('Documentation Generation Tool (PowerDocu)',
     'What is Cydef?',
     'One clause. The row names them as the collaborator and never says who they are.', 2),
    ('SAGE Copilot AI',
     'What are SAGE, ME5 and the Delta Agent, in plain terms?',
     'Three acronyms in a single sentence, none expanded. Say what each one is and what the integration does.', 4),
    ('SAGE Copilot AI',
     'What does "replicate GPT-4.1 concepts" mean, and which parts did you build versus SAGE supplying?',
     'As written it could mean anything from prompt design to model fine-tuning, and the ownership split is '
     'unclear.', 4),
    ('1-Day Power Platform Bootcamp',
     'What is 815 SQN, and where does the 35-attendee figure come from?',
     'The squadron number means nothing outside the RSAF, and the headcount has no source.', 3),
    ('National AI Student Challenge 2025',
     'Did the entry place, and what did the judges assess?',
     'The certificate is already shown, so this is about what sat behind it.', 3),
    ('Project Management Tracker',
     'Which parts of it did you build?',
     'The row says it is in use across RSAF teams and scaling to enterprise, but not what your hand in it was. '
     'You are separately named as its Security Reviewer, so if that is the whole of your involvement, say that '
     'and the row gets rewritten around it.', 3),
]

n = 14
for where, ask, hint, lines in log:
    question(n, 'projects.html  ·  ' + where, ask, hint, lines=lines)
    n += 1

# ─────────────────── closing ───────────────────
doc.add_page_break()
para('One last thing, and it is not a question', size=13, bold=True, color=INK, after=4)
para('The em dashes are gone from every page of the site, including the meta descriptions that show up in '
     'search results and link previews. Two places still have them, on purpose:', size=10, color=BODY, after=4)
for b in [
    'Inside the rebuilt app screens on the project log. The BOLDFACE quiz lines, GRID\'s upload hint, its '
    'planning-view notice and its step banners are the apps\' own strings, reproduced exactly. The page claims '
    'those screens are rebuilt from the real interface code, so editing their wording would make that claim '
    'false. Say the word and they get changed anyway.',
    'In the source code comments, which no visitor ever sees.',
]:
    p = doc.add_paragraph(style='List Bullet')
    r = p.add_run(b); r.font.size = Pt(10); r.font.color.rgb = BODY
    p.paragraph_format.space_after = Pt(4)

out = os.path.join('docs', 'portfolio-v3-content-questionnaire.docx')
doc.save(out)
print('wrote', out, os.path.getsize(out), 'bytes')
